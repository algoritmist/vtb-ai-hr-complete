# analyzer.py

from sentence_transformers import SentenceTransformer, util
import torch
import os
import re
from pathlib import Path
from striprtf.striprtf import rtf_to_text
from docx import Document
from pypdf import PdfReader
import json
from typing import Union, List, Dict, Optional

# ==============================
# 1. ИЗВЛЕЧЕНИЕ ТЕКСТА ИЗ ФАЙЛОВ
# ==============================


def extract_text_as_single_line(file_path: str) -> str:
    """Извлекает и очищает текст из .docx, .pdf, .rtf"""

    def clean_special_chars(text: str) -> str:
        if not text:
            return ""
        text = text.replace('\\t', ' ').replace('\t', ' ')
        text = re.sub(
            r'[\n\r\f\v\u00a0\u1680\u2000-\u200F\u2028-\u202F\u205F\u2060\u3000]', ' ', text)
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', text)
        text = re.sub(r'[\u200E\u200F\u202A-\u202E]', '', text)
        text = re.sub(r'[•▪▶➢\*\•\‣\⁃\-\•]', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def extract_from_docx(filepath: str) -> str:
        try:
            doc = Document(filepath)
            paragraphs = [p.text.strip()
                          for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return ' '.join(paragraphs)
        except Exception as e:
            print(f"Ошибка DOCX: {str(e)}")
            return ""

    def extract_from_pdf(filepath: str) -> str:
        try:
            text_parts = []
            with open(filepath, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return ' '.join(text_parts)
        except Exception as e:
            print(f"Ошибка PDF: {str(e)}")
            return ""

    def extract_from_rtf(filepath: str) -> str:
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as file:
                rtf_content = file.read()
            return rtf_to_text(rtf_content)
        except Exception as e:
            print(f"Ошибка RTF: {str(e)}")
            return ""

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    extension = Path(file_path).suffix.lower()

    if extension == '.rtf':
        text = extract_from_rtf(file_path)
    elif extension == '.docx':
        text = extract_from_docx(file_path)
    elif extension == '.pdf':
        text = extract_from_pdf(file_path)
    else:
        raise ValueError(f"Неподдерживаемый формат: {extension}")

    return clean_special_chars(text) if text else ""

# ==============================
# 2. ПАРСИНГ ТЕКСТА ВАКАНСИИ
# ==============================


def parse_text_to_dict(text: str) -> Dict[str, Union[str, List[str]]]:
    known_fields = [
        'Наименование поля', 'Значение', 'Статус', 'Название', 'Регион', 'Город',
        'Адрес', 'Тип трудового', 'Тип занятости', 'Текст график работы',
        'Доход (руб/мес)', 'Оклад макс. (руб/мес)', 'Оклад мин. (руб/мес)',
        'Годовая премия (%)', 'Тип премирования. Описание', 'Обязанности (для публикации)',
        'Требования (для публикации)', 'Будет преимуществом:', 'Уровень образования',
        'Требуемый опыт работы', 'Знание специальных программ', 'Навыки работы на компьютере',
        'Знание иностранных языков', 'Уровень владения языка', 'Наличие командировок',
        'Дополнительная информация'
    ]

    field_pattern = '|'.join(re.escape(field) for field in known_fields)
    parts = re.split(f'({field_pattern})', text)
    parts = [part.strip() for part in parts if part and part.strip()]

    result = {}
    i = 0
    while i < len(parts):
        part = parts[i]
        if part in known_fields:
            current_field = part
            i += 1
            value_parts = []
            while i < len(parts) and parts[i] not in known_fields:
                value_parts.append(parts[i])
                i += 1
            value = ' '.join(value_parts).strip()
            if any(keyword in current_field.lower() for keyword in ['обязанности', 'требования', 'преимуществом']):
                if ';' in value:
                    value_list = [item.strip()
                                  for item in value.split(';') if item.strip()]
                    result[current_field] = value_list
                else:
                    lines = re.split(r'\n\s*\n|\n(?=\d+\.|\•|\-)', value)
                    if len(lines) > 1:
                        result[current_field] = [line.strip()
                                                 for line in lines if line.strip()]
                    else:
                        result[current_field] = value
            else:
                result[current_field] = value
        else:
            i += 1
    return result


def clean_and_format_dict(data_dict: Dict) -> Dict:
    cleaned_dict = {}
    for key, value in data_dict.items():
        clean_key = key.strip()
        if isinstance(value, str):
            clean_value = re.sub(r'\s+', ' ', value.strip())
        else:
            clean_value = value
        cleaned_dict[clean_key] = clean_value
    return cleaned_dict


def parse_vacancy_from_json(jsonchick: Dict) -> Dict:
    raw_data = jsonchick
    vacancy = {
        "title": raw_data.get("Название", "").strip(),
        "location": (raw_data.get("Город", "") or raw_data.get("Регион", "")).strip(),
        "responsibilities": [],
        "requirements": [],
        "preferred": [],
        "education": raw_data.get("Уровень образования", "").strip(),
        "experience_years": raw_data.get("Требуемый опыт работы", "").strip(),
        "travel": raw_data.get("Наличие командировок", "").strip(),
    }

    def process_raw(raw):
        if isinstance(raw, list):
            return [item.strip() for item in raw if item.strip()]
        elif isinstance(raw, str) and raw.strip():
            sentences = [s.strip() for s in raw.split('.') if s.strip()]
            if len(sentences) == 1 and '\n' in raw:
                sentences = [s.strip() for s in raw.split('\n') if s.strip()]
            return sentences
        return []

    vacancy["responsibilities"] = process_raw(
        raw_data.get("Обязанности (для публикации)", ""))
    vacancy["requirements"] = process_raw(
        raw_data.get("Требования (для публикации)", ""))
    vacancy["preferred"] = process_raw(
        raw_data.get("Будет преимуществом:", ""))

    return vacancy

# ==============================
# 3. АНАЛИЗАТОР СХОЖЕСТИ
# ==============================


class InterviewAnalyzer:
    def __init__(self, model_name='ai-forever/sbert_large_nlu_ru', device=None, threshold=0.5, default_soft_skill_score=0.3):
        self.device = device if device else (
            "cuda" if torch.cuda.is_available() else "cpu")
        self.model = SentenceTransformer(model_name, device=self.device)
        self.threshold = threshold
        self.default_soft_skill_score = default_soft_skill_score
        self.CATEGORIES_CONFIG = self._get_categories_config()

    def _get_categories_config(self):
        return {
            "technical_skills": [
                "настройка", "оборудование", "сервер", "сеть", "raid", "массив дисков", "восстановление дисков",
                "bmc", "bios", "python", "sql", "cisco", "mikrotik", "ssh", "ubuntu", "windows server", "скрипт",
                "кабель", "монтаж", "демонтаж", "техобслуживание", "сборка", "диагностика", "инцидент", "подключение",
                "схд", "коммутатор", "firewall", "пк", "сетевое", "linux", "windows", "api", "cli", "bash", "powershell",
                "html", "css", "rest", "graphql", "docker", "kubernetes", "базы данных", "orm", "json", "xml",
                "отказоустойчивость", "резервирование", "восстановление", "логи", "мониторинг", "отказ диска",
                "javascript", "js", "typescript", "react", "vue", "angular", "фронтенд", "spa", "верстка", "redux",
                "node.js", "django", "flask", "spring", "backend", "бэкенд", "микросервисы", "java", "c#", "go",
                "devops", "ansible", "terraform", "jenkins", "ci/cd", "k8s", "aws", "azure", "gcp", "nginx",
                "бизнес-аналитик", "риск-менеджер", "fraud", "антифрод", "кредитный аналитик", "scoring",
                "системный аналитик", "специалист цод", "цод", "дата-центр", "rack", "х86", "dcim"
            ],
            "communication_skills": [
                "речь", "коммуникация", "обучение", "консультация", "взаимодействие", "координация",
                "общение", "отчет", "подготовка отчетов", "сопровождение", "клиент", "консультирование",
                "презентация", "переговоры", "деловая переписка", "грамотный", "объяснил", "договорился"
            ],
            "case_projects": [
                "проект", "внедрение", "разработка", "тест", "анализ", "оптимизация", "реализация",
                "автоматизация", "восстановительные работы", "mvp", "релиз", "интеграция", "улучшил",
                "сократил", "увеличил", "добился", "результат", "метрика", "kpi", "экономия"
            ]
        }

    def categorize_item(self, item_text: str) -> str:
        item_lower = item_text.lower()
        for category, keywords in self.CATEGORIES_CONFIG.items():
            if any(kw in item_lower for kw in keywords):
                return category
        return "experience_relevance"

    def extract_experience_from_text(self, text_list: List[str]) -> int:
        full_text = " ".join(text_list).lower()
        patterns = [
            r'(\d+)\s*лет?\s*(\d+)?\s*месяц[а-я]*',
            r'(\d+)\s*года?\s*(\d+)?\s*месяц[а-я]*',
            r'(\d+)\s*лет?',
            r'(\d+)\s*года?',
            r'(\d+)\s*год[а-я]*',
            r'(\d+)\s*месяц[а-я]*'
        ]
        total_months = 0
        for pattern in patterns:
            matches = re.findall(pattern, full_text)
            for match in matches:
                if isinstance(match, tuple):
                    years = int(match[0]) if match[0] else 0
                    months = int(match[1]) if len(
                        match) > 1 and match[1] else 0
                else:
                    num = int(match)
                    if "лет" in pattern or "год" in pattern:
                        years, months = num, 0
                    else:
                        years, months = 0, num
                total_months += years * 12 + months
        return total_months

    def parse_required_experience(self, exp_str: str) -> tuple:
        if not exp_str:
            return 0, 0
        s = exp_str.lower().strip()
        if any(phrase in s for phrase in ["не требуется", "без опыта", "нет требований"]):
            return 0, 0
        s = re.sub(r'[–—\-\s]+', '-', s)
        if "от" in s:
            match = re.search(
                r'от[\s\-]*(\d+)[\s\-]*(лет|года?|месяц[а-я]*)', s)
            if match:
                num = int(match.group(1))
                unit = match.group(2)
                months = num * 12 if 'лет' in unit or 'год' in unit else num
                return months, 999
        if "более" in s:
            match = re.search(r'более[\s\-]*(\d+)[\s\-]*(лет|года?)', s)
            if match:
                num = int(match.group(1))
                return num * 12, 999
        match = re.search(r'(\d+)[\-](\d+)[\s\-]*(лет|года?)', s)
        if match:
            min_y = int(match.group(1))
            max_y = int(match.group(2))
            return min_y * 12, max_y * 12
        match = re.search(r'(\d+)[\s\-]*(лет|года?|месяц[а-я]*)', s)
        if match:
            num = int(match.group(1))
            unit = match.group(2)
            months = num * 12 if 'лет' in unit or 'год' in unit else num
            return months, months
        return 0, 0

    def match_experience(self, candidate_months: int, vacancy_exp_str: str) -> float:
        min_req, max_req = self.parse_required_experience(vacancy_exp_str)
        if max_req == 0 and min_req == 0:
            return 1.0
        if min_req <= candidate_months <= max_req:
            return 1.0
        elif candidate_months < min_req:
            return max(0.0, 1.0 - (min_req - candidate_months) / (min_req or 1))
        elif candidate_months > max_req and max_req < 999:
            return 0.7
        elif max_req == 999:
            return 1.0 if candidate_months >= min_req else max(0.0, 1.0 - (min_req - candidate_months) / min_req)
        else:
            return 0.0

    def evaluate_answer_depth(self, answer_text: str) -> Dict:
        text = answer_text.lower()
        indicators = {
            "example": ["на проекте", "в компании", "у нас был", "когда я работал", "однажды"],
            "detail": ["использовал", "настроил", "применил", "инструмент", "технология", "версия", "v1", "v2"],
            "result": ["сократил", "увеличил", "добился", "улучшил", "экономия", "результат", "kpi"]
        }
        score = 0
        reasons = []
        for key, phrases in indicators.items():
            if any(phrase in text for phrase in phrases):
                score += 1
                reasons.append(key)
        depth_level = "surface"
        if score >= 3:
            depth_level = "detailed"
        elif score >= 1:
            depth_level = "example"
        return {
            "depth_level": depth_level,
            "score": score,
            "indicators_found": reasons
        }

    def match_text_to_requirement(self, source_text: str, requirement_text: str) -> tuple:
        if not source_text.strip() or not requirement_text.strip():
            return False, None, 0.0
        req_prompt = f"Требование: {requirement_text}"
        src_prompt = f"Текст кандидата: {source_text}"
        req_emb = self.model.encode(
            req_prompt, convert_to_tensor=True, device=self.device)
        src_emb = self.model.encode(
            src_prompt, convert_to_tensor=True, device=self.device)
        score = util.cos_sim(req_emb, src_emb).item()
        is_matched = score >= self.threshold
        return is_matched, source_text, score

    def analyze(self, resume_input: Union[str, List[str]], vacancy: Dict, weights: Optional[Dict] = None, return_features: bool = False) -> Dict:
        base_weights = {
            "technical_skills": 0.4,
            "experience_years_match": 0.3,
            "communication_skills": 0.15,
            "case_projects": 0.1,
            "experience_relevance": 0.05
        }
        if weights is None:
            weights = base_weights.copy()

        # Определяем тип входных данных
        if isinstance(resume_input, str):
            fragments = self._parse_resume_into_fragments(resume_input)
            candidate_total_months = self.extract_experience_from_text([
                                                                       resume_input])
            all_source_texts = fragments
            is_interview = False
        elif isinstance(resume_input, list) and all(isinstance(x, str) for x in resume_input):
            answers_text_list = [
                ans.strip() for ans in resume_input if isinstance(ans, str) and ans.strip()]
            candidate_total_months = self.extract_experience_from_text(
                answers_text_list)
            all_source_texts = answers_text_list
            is_interview = True
        else:
            raise ValueError("resume_input должен быть str или list[str]")

        required_exp_str = vacancy.get(
            "Требуемый опыт работы", "") or vacancy.get("experience_years", "")
        exp_match_score = self.match_experience(
            candidate_total_months, required_exp_str)

        all_vacancy_items = (
            [{"text": r, "category": "responsibilities"} for r in vacancy.get("responsibilities", []) if r] +
            [{"text": r, "category": "requirements"} for r in vacancy.get("requirements", []) if r] +
            [{"text": r, "category": "preferred"}
                for r in vacancy.get("preferred", []) if r]
        )

        present_categories = set()
        for item in all_vacancy_items:
            cat = self.categorize_item(item["text"])
            if cat in weights:
                present_categories.add(cat)

        if not required_exp_str.strip() or self.parse_required_experience(required_exp_str) == (0, 0):
            present_categories.discard("experience_years_match")
        else:
            present_categories.add("experience_years_match")

        active_weights = {k: v for k,
                          v in weights.items() if k in present_categories}
        if not active_weights:
            active_weights = {"experience_relevance": 1.0}

        total_weight = sum(active_weights.values())
        if total_weight > 0:
            active_weights = {k: v / total_weight for k,
                              v in active_weights.items()}

        category_scores = {cat: []
                           for cat in active_weights if cat != "experience_years_match"}
        matched_items = []

        for item in all_vacancy_items:
            best_score = 0.0
            best_source = None
            best_depth = None

            for src_text in all_source_texts:
                is_matched, source, score = self.match_text_to_requirement(
                    src_text, item["text"])
                if score > best_score:
                    best_score = score
                    best_source = source
                    if is_interview:
                        best_depth = self.evaluate_answer_depth(source)

            cat = self.categorize_item(item["text"])

            matched_item = {
                "item": item["text"],
                "found": bool(best_score >= self.threshold),
                "source": best_source if best_source else None,
                "similarity_score": round(best_score, 3),
                "category": cat
            }

            if best_depth and is_interview:
                matched_item["depth_analysis"] = best_depth

            matched_items.append(matched_item)

            if cat in category_scores:
                category_scores[cat].append(best_score)

        criteria_scores = {}

        for cat, scores in category_scores.items():
            if scores:
                avg = sum(scores) / len(scores)
            else:
                cat_items = [item for item in all_vacancy_items if self.categorize_item(
                    item["text"]) == cat]
                if cat_items:
                    if cat == "communication_skills":
                        avg = self.default_soft_skill_score
                    else:
                        avg = 0.0
                else:
                    avg = 0.0
            criteria_scores[cat] = round(avg * 100, 1)

        if "experience_years_match" in active_weights:
            criteria_scores["experience_years_match"] = round(
                exp_match_score * 100, 1)

        total = sum(criteria_scores.get(cat, 0) *
                    active_weights[cat] for cat in active_weights)
        total_match_percent = round(total, 1)

        result = {
            "total_match_percent": total_match_percent,
            "criteria_scores": criteria_scores,
            "matched_items": matched_items,
            "candidate_experience": {
                "total_months": int(candidate_total_months),
                "total_years": round(candidate_total_months / 12, 1),
                "required_experience": required_exp_str,
                "match_score": round(exp_match_score, 3)
            },
            "weights_used": active_weights
        }

        if return_features:
            result["features_used"] = self.extract_features_from_vacancy(
                vacancy)

        return result

    def extract_features_from_vacancy(self, vacancy: Dict) -> List[Dict]:
        all_vacancy_items = (
            [{"text": r, "category": "responsibilities"} for r in vacancy.get("responsibilities", []) if r] +
            [{"text": r, "category": "requirements"} for r in vacancy.get("requirements", []) if r] +
            [{"text": r, "category": "preferred"}
                for r in vacancy.get("preferred", []) if r]
        )
        features = []
        for item in all_vacancy_items:
            cat = self.categorize_item(item["text"])
            if cat in ["technical_skills", "communication_skills", "case_projects"]:
                features.append({
                    "requirement": item["text"],
                    "category": cat
                })
        return features

    def _parse_resume_into_fragments(self, resume_text: str) -> List[str]:
        try:
            import razdel
        except ImportError:
            sentences = []
            for part in re.split(r'[.!?]+(?=\s+[А-Я]|$)', resume_text):
                for line in part.split('\n'):
                    stripped = line.strip()
                    if len(stripped) > 5:
                        sentences.append(stripped)
            return sentences

        text = re.sub(r'(Опыт работы\s*—[^—]+)\s*\1+',
                      r'\1', resume_text, flags=re.IGNORECASE)
        text = re.sub(r'(Образование\s+)+', 'Образование ',
                      text, flags=re.IGNORECASE)
        text = re.sub(r'(Навыки\s+)+', 'Навыки ', text, flags=re.IGNORECASE)
        text = re.sub(r'(Дополнительная информация\s+)+',
                      'Дополнительная информация ', text, flags=re.IGNORECASE)

        parts = re.split(
            r'(?=\n*(Опыт работы|Образование|Навыки|Дополнительная информация|Обо мне)\s*[:—]?)',
            text,
            flags=re.IGNORECASE
        )

        fragments = []
        i = 0
        while i < len(parts):
            part = parts[i].strip()
            if not part or len(part) < 3:
                i += 1
                continue

            if re.search(r'Опыт работы', part, re.IGNORECASE):
                i += 1
                if i < len(parts):
                    content = parts[i].strip()
                    exp_blocks = self._extract_experience_blocks(content)
                    for block in exp_blocks:
                        sentences = [sent.text.strip() for sent in razdel.sentenize(
                            block) if len(sent.text.strip()) > 2]
                        fragments.extend(sentences)
            elif re.search(r'Навыки', part, re.IGNORECASE):
                i += 1
                if i < len(parts):
                    skills_text = parts[i].strip()
                    skills_list = re.split(r'[,;\n•–—\-\s]\s*', skills_text)
                    for skill in skills_list:
                        skill_clean = skill.strip(" .,;:-–—")
                        if len(skill_clean) > 2:
                            fragments.append(skill_clean)
            elif re.search(r'Образование', part, re.IGNORECASE):
                i += 1
                if i < len(parts):
                    edu_text = parts[i].strip()
                    sentences = [sent.text.strip() for sent in razdel.sentenize(
                        edu_text) if len(sent.text.strip()) > 2]
                    fragments.extend(sentences)
            elif re.search(r'Дополнительная информация|Обо мне', part, re.IGNORECASE):
                i += 1
                if i < len(parts):
                    about_text = parts[i].strip()
                    sentences = [sent.text.strip() for sent in razdel.sentenize(
                        about_text) if len(sent.text.strip()) > 2]
                    fragments.extend(sentences)
            else:
                sentences = [sent.text.strip() for sent in razdel.sentenize(
                    part) if len(sent.text.strip()) > 2]
                fragments.extend(sentences)
            i += 1

        seen = set()
        unique_fragments = []
        for frag in fragments:
            if frag and len(frag) > 2 and frag not in seen:
                seen.add(frag)
                unique_fragments.append(frag)

        return unique_fragments

    def _extract_experience_blocks(self, text: str) -> List[str]:
        if not text:
            return []
        pattern = r'([А-Яа-я]+\s+\d{4}\s*—\s*[А-Яа-я\s\d]+)'
        positions = [match.start() for match in re.finditer(pattern, text)]
        if not positions:
            return [text.strip()]
        blocks = []
        for i in range(len(positions)):
            start = positions[i]
            end = positions[i + 1] if i + 1 < len(positions) else len(text)
            block = text[start:end].strip()
            if len(block) > 10:
                blocks.append(block)
        return blocks

# ==============================
# 4. УДОБНЫЕ ИНТЕРФЕЙСЫ
# ==============================


def analyze_vacancy_vs_resume(vacancy_file: str, resume_file: str) -> Dict:
    """Анализирует схожесть вакансии и резюме"""
    vacancy_text = extract_text_as_single_line(vacancy_file)
    resume_text = extract_text_as_single_line(resume_file)

    vacancy_dict_raw = parse_text_to_dict(vacancy_text)
    vacancy_dict_clean = clean_and_format_dict(vacancy_dict_raw)
    vacancy_structured = parse_vacancy_from_json(vacancy_dict_clean)

    analyzer = InterviewAnalyzer()
    result = analyzer.analyze(
        resume_text, vacancy_structured, return_features=True)
    return result


def analyze_vacancy_vs_interview(vacancy_text: str, interview_answers: List[str]) -> Dict:
    """Анализирует схожесть вакансии и ответов на интервью"""
    vacancy_dict_raw = parse_text_to_dict(vacancy_text)
    vacancy_dict_clean = clean_and_format_dict(vacancy_dict_raw)
    vacancy_structured = parse_vacancy_from_json(vacancy_dict_clean)

    analyzer = InterviewAnalyzer()
    result = analyzer.analyze(
        interview_answers, vacancy_structured, return_features=True)
    return result

# ==============================
# 5. LLMAnalyzer
# ==============================


class LLMAnalyzer:
    """Заглушка для замены старого LLMAnalyzer — теперь использует InterviewAnalyzer"""

    def __init__(self, api_key: str, db_api_url: Optional[str] = None):
        self.api_key = api_key
        self.db_api_url = db_api_url
        self.analyzer = InterviewAnalyzer()

    def analyze_text(self, vacancy_text: str, history_text: str) -> str:
        """
        Принимает JSON-строку с ответами кандидата, возвращает результат анализа.
        Пример history_text: '["Писал скрипты на Python", "Использовал Docker"]'
        """
        try:
            interview_answers = json.loads(history_text)
            if not isinstance(interview_answers, list):
                raise ValueError("history_text должен быть списком строк")
            result = analyze_vacancy_vs_interview(
                vacancy_text, interview_answers)
            return json.dumps(result, ensure_ascii=False, indent=2)
        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False, indent=2)

    def analyze_resume(self, vacancy_file: str, resume_file: str) -> str:
        """Анализ резюме против вакансии"""
        try:
            result = analyze_vacancy_vs_resume(vacancy_file, resume_file)
            return json.dumps(result, ensure_ascii=False, indent=2)
        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False, indent=2)
